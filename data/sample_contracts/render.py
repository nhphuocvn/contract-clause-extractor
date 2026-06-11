"""Render the synthetic deal package to DOCX and PDF.

Both renderers consume the same canonical text modules (`_doc_a_content.py`,
`_doc_b_content.py`) so DOCX and PDF cannot drift.

Usage:
    python -m data.sample_contracts.render

Writes four files into `data/sample_contracts/`:
    gpu_purchase_agreement.docx
    gpu_purchase_agreement.pdf
    warrant_agreement.docx
    warrant_agreement.pdf

Section headings are rendered as "<N>. <ALL-CAPS TITLE>" on their own line so
that the existing chunker in `index.SECTION_HEADER` recognizes them when the
PDFs are read back through PyPDF2.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from data.sample_contracts import _doc_a_content as doc_a
from data.sample_contracts import _doc_b_content as doc_b

OUT_DIR = Path(__file__).resolve().parent


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------


def _docx_set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def render_docx(content_module, out_path: Path) -> None:
    document = Document()
    _docx_set_default_font(document)
    for section in document.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Title
    title = document.add_paragraph()
    title_run = title.add_run(content_module.TITLE)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = 1  # center

    # Notice
    notice = document.add_paragraph()
    notice_run = notice.add_run(content_module.NOTICE)
    notice_run.italic = True
    notice_run.font.size = Pt(9)
    notice.alignment = 1  # center

    document.add_paragraph()  # spacer

    # Preamble
    document.add_paragraph(content_module.PREAMBLE)

    # Recitals (Doc A only)
    if hasattr(content_module, "RECITALS"):
        document.add_paragraph(content_module.RECITALS)

    document.add_paragraph()  # spacer

    # Sections
    for number, heading, paragraphs in content_module.SECTIONS:
        h = document.add_paragraph()
        h_run = h.add_run(f"{number}. {heading}")
        h_run.bold = True
        h_run.font.size = Pt(12)
        for para in paragraphs:
            document.add_paragraph(para)

    document.add_paragraph()  # spacer
    document.add_paragraph(content_module.SIGNATURE_BLOCK)

    document.save(str(out_path))


# ---------------------------------------------------------------------------
# PDF renderer
# ---------------------------------------------------------------------------


def _pdf_styles() -> dict:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "DealTitle",
            parent=base["Title"],
            fontName="Times-Bold",
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "notice": ParagraphStyle(
            "DealNotice",
            parent=base["Italic"],
            fontName="Times-Italic",
            fontSize=9,
            alignment=TA_CENTER,
            textColor="#666666",
            spaceAfter=24,
        ),
        # Section heading. spaceBefore ensures it lands on its own line in
        # extracted text; the "<N>. ALL CAPS" form matches index.SECTION_HEADER.
        "heading": ParagraphStyle(
            "DealHeading",
            parent=base["Heading2"],
            fontName="Times-Bold",
            fontSize=12,
            spaceBefore=14,
            spaceAfter=6,
            keepWithNext=True,
            alignment=TA_LEFT,
        ),
        "body": ParagraphStyle(
            "DealBody",
            parent=base["BodyText"],
            fontName="Times-Roman",
            fontSize=12,
            alignment=TA_JUSTIFY,
            spaceAfter=10,
            leading=16,
        ),
        "preamble": ParagraphStyle(
            "DealPreamble",
            parent=base["BodyText"],
            fontName="Times-Roman",
            fontSize=12,
            alignment=TA_JUSTIFY,
            spaceAfter=14,
            leading=16,
        ),
        "signature": ParagraphStyle(
            "DealSignature",
            parent=base["BodyText"],
            fontName="Times-Roman",
            fontSize=11,
            alignment=TA_LEFT,
            spaceBefore=24,
            leading=14,
        ),
    }


def render_pdf(content_module, out_path: Path) -> None:
    styles = _pdf_styles()
    document = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=content_module.TITLE,
    )

    flow: list = []
    flow.append(Paragraph(content_module.TITLE, styles["title"]))
    flow.append(Paragraph(content_module.NOTICE, styles["notice"]))
    flow.append(Paragraph(content_module.PREAMBLE, styles["preamble"]))
    if hasattr(content_module, "RECITALS"):
        flow.append(Paragraph(content_module.RECITALS, styles["preamble"]))
    flow.append(Spacer(1, 8))

    for number, heading, paragraphs in content_module.SECTIONS:
        flow.append(Paragraph(f"{number}. {heading}", styles["heading"]))
        for para in paragraphs:
            flow.append(Paragraph(para, styles["body"]))

    # Signature block: split on newlines so reportlab renders one line each.
    for line in content_module.SIGNATURE_BLOCK.split("\n"):
        if line.strip() == "":
            flow.append(Spacer(1, 6))
        else:
            flow.append(Paragraph(line, styles["signature"]))

    document.build(flow)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


JOBS = [
    (doc_a, "gpu_purchase_agreement"),
    (doc_b, "warrant_agreement"),
]


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for module, stem in JOBS:
        docx_path = OUT_DIR / f"{stem}.docx"
        pdf_path = OUT_DIR / f"{stem}.pdf"
        render_docx(module, docx_path)
        render_pdf(module, pdf_path)
        print(f"wrote {docx_path}")
        print(f"wrote {pdf_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
